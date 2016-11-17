#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
docx2bb:
Create BlackBoard (bb) test questions (text) import file from a Microsoft Word *.docx document. 
Supported question types are: True/False, Multiple choice, Matching, Essay, and (simple) Fill 
in the blank. ExamFormat-Sample.docx shows a sample exam format for use with docx2bb. 
Unicode-to-ASCII replacement rules from 'docx2bb.json' data file can be optionally applied.

Syntax:
    python docx2bb.py [options] [docx_filename]
options:
	--verbose | -v display verbose messages
	--help | -h display help message

Disclaimer:
docx2bb is provided with no warranties, use it if you find it useful. docx2bb is designed to 
keep your *.docx document unchanged, but the author assumes no liabilities from use of 
this tool, including if it eats your exam :).

Code by Sinan Salman, 2016

Version History:
27.10.16	0.10	inital release on bitbucket
29.10.16	0.11	fix outline level identification issue
07.11.16	0.12	add PageBreak elimination logic, fix line id reporting in verbose, and made T/F RegEx substitution case insensitive
17.11.16	0.13	add default values for unicode2ascii replacements even if docx2bb.jason file is absent
"""

HELP_MSG = """Options:
	--verbose | -v  display verbose messages
	--help    | -h  display help message
"""

INSTALL_MSG = """
OSX and Linux:
	pip install python-docx

Windows platforms:
	Download lxml library wheel:
		http://www.lfd.uci.edu/~gohlke/pythonlibs/#lxml
	Install lxml library:
		pip install [downloadedfile]
	Install python-docx library:
		pip install python-docx
"""
		
__app__ = 		"docx2bb.py"
__author__ = 	"Sinan Salman (sinan.salman[at]gmail.com)"
__version__ = 	"v0.13"
__date__ = 		"Nov 17, 2016"
__copyright__ = "Copyright (c)2016 Sinan Salman"
__license__ = 	"GPLv3"
__website__	=	"https://bitbucket.org/sinansalman/docx2bb"

### Initialization #######################################################################

import re
import os
import sys
import platform
        
verbose = False
unicode2ascii = {'rules':{'“':'"','”':'"','‘':"'",'’':"'",'–':'-'}, 
				 'notallowed':'[^a-zA-Z0-9 §±!@#$%^&*()\\-_=+[\\]{};:\'\"\\\\|<>,./?`~\\n]'}
WordFileName = ""
data = []
BBtext = ""
TextWidth = 0
QuestionTypes = {'T/F':0, 'M/C':0,'MAT':0,'FIB':0,'ESSAY':0,'Warning':0}

# encoding=utf8  ==> fix for Python 2.7
if sys.version_info[0] == 2:
	reload(sys)  
	sys.setdefaultencoding('utf8')

# platform identification for correct output in Windows
if platform.system() == 'Windows':
        class bcolors:
                WARNING = 	''
                FAIL = 		''
                ENDC = 		''
else:
        class bcolors:
                WARNING = 	'\033[33m'
                FAIL = 		'\033[31m'
                ENDC = 		'\033[0m'

### Script Management ####################################################################

def ProcessCLI():
	"""Process CLI parameters"""
	global verbose
	global WordFileName
	global unicode2ascii
	global TextWidth
	
	# Get terminal width
	try:
		TextWidth = os.get_terminal_size()[0]
	except:
		TextWidth = 80

	print ('{:} | {:} | {:} | {:} License\nDownload latest version at {:}\n'.format(__app__,__version__,__date__,__license__,__website__))
	
	# handle arguments and load settings from JSON file
	if len(sys.argv) == 1:
		print ("Syntax:\n\tpython docx2bb.py [options] [docx_filename]")
		print (HELP_MSG)
		print_Fail ("Missing argument")
	if '--verbose' in sys.argv or '-v' in sys.argv:
		print ("*** Option: verbose mode")
		verbose = True
	if '--help' in sys.argv or '-h' in sys.argv:
		print("Syntax:\n\tpython docx2bb.py [options] [docx_filename]")
		print(HELP_MSG)
		sys.exit(0)
	WordFileName = sys.argv[-1]
	if not os.path.isfile(WordFileName):
		print_Fail ("can't find file: {:}. Make sure [docx_filename] is the last argument.".format(WordFileName))
	import json
	if os.path.isfile('docx2bb.json'):
		with open('docx2bb.json',encoding='utf-8') as jsonfile:
			unicode2ascii = json.load(jsonfile)
			if verbose:
				print ("*** Option: using docx2bb.json file")
	
### Analyze Document and Convert to BB Text Format #######################################

def RunScript():
	"""Process word file and create BB text import file"""

	# try to import python-docx module
	try:
		import docx
	except ImportError as e:
		print ("Please install python-docx library first using:")
		print (INSTALL_MSG)
		sys.exit(0)

	# open docx file
	if verbose: print ('Reading Docx file...')
	doc = docx.Document(WordFileName)

	# extract data from docx file
	global Data
	n=0
	if verbose: print ("Found {:} paragraphs, parsing and converting unicode to ascii...".format(len(doc.paragraphs)))
	for p in doc.paragraphs:
		n += 1
		temp={'No':0,'text':u2a(p.text),'outline':0,'allBold':False,'trueBold':False,'falseBold':False,'list':False}
		bold = True
		for r in p.runs:
			if r.bold == None: bold = False
			if r.bold == True and r.text.strip(' ').lower() == 'true': 	temp['trueBold'] = True
			if r.bold == True and r.text.strip(' ').lower() == 'false': temp['falseBold'] = True
		temp['allBold'] = bold
		temp['list'], temp['outline'] = GetListOutline(p)
		temp['Q'] = 0
		temp['No'] = n
		data.append(temp)
		
	# print data object for debugging
	if verbose: 
		print ('\nBefore clean up:')
		print ('    out   is') 
		print (' #  line  list text') 
		print ('~~~ ~~~~ ~~~~~ ~~~~') 
		i = 1
		for d in data:
			print ("{:3} {:4} {:5} {:}".format(d['No'],d['outline'],str(d['list']),d['text'][:(TextWidth-15)])) 
			i += 1

	if verbose:	print ('\nClean up:')	
	# delete empty paragraphs
	empty_p = []
	for i in range(len(data)):
		if data[i]['text'].strip(' ') == '':
			empty_p.append(i)
	if empty_p != []:
		if verbose:
			print ('Removing {:} empty line(s): {:}'.format(len(empty_p),[data[x]['No'] for x in empty_p]))	
		for i in reversed(range(len(empty_p))): # go backward to delete lines without missing up the index
			del data[empty_p[i]]

	# delete pagebreaks
	PageBreak = []
	for i in range(len(data)):
		if re.search('^\s*\n+$',data[i]['text']):
			PageBreak.append(i)
	if PageBreak != []:
		if verbose:
			print ('Removing {:} pagebreak(s) at: {:}'.format(len(PageBreak),[data[x]['No'] for x in PageBreak]))	
		for i in reversed(range(len(PageBreak))): # go backward to delete lines without missing up the index
			del data[PageBreak[i]]

	# delete non-list paragraphs
	nonlist_p = []
	for i in range(len(data)):
		if data[i]['list'] == False:
			nonlist_p.append(i)
	if nonlist_p != []:
		if verbose:
			print ('Removing {:} none-list line(s): {:}'.format(len(nonlist_p),[data[x]['No'] for x in nonlist_p]))	
		for i in reversed(range(len(nonlist_p))): # go backward to delete lines without missing up the index
			del data[nonlist_p[i]]

	# identify question start/end paragraph positions
	Qbeg_pos = [0]
	Qend_pos = []
	Qid = 1
	data[0]['Q'] = Qid
	for i in range(1, len(data)-1):
		if data[Qbeg_pos[-1]]['outline'] < data[i]['outline']:
			data[i]['Q'] = Qid
		else:
			Qend_pos.append(i-1)
			Qbeg_pos.append(i)
			Qid += 1
			data[i]['Q'] = Qid
	Qend_pos.append(len(data)-1) 
	data[-1]['Q'] = Qid

	# print data object for debugging
	if verbose: 
		print ('\nAfter clean up:')
		print ('        out   all  true  false is') 
		print (' #   Q  line  Bold Bold  Bold  list  text') 
		print ('~~~ ~~~ ~~~~ ~~~~~ ~~~~~ ~~~~~ ~~~~~ ~~~~') 
		i = 1
		for d in data:
			print ("{:3} {:3} {:4} {:5} {:5} {:5} {:5} {:}".format(d['No'],d['Q'],d['outline'],str(d['allBold']),str(d['trueBold']),str(d['falseBold']),str(d['list']),d['text'][:(TextWidth-37)])) 
			i += 1
	if len(Qbeg_pos) != len(Qend_pos): print_Fail ("problem in parsing question lines.\n\tQ_begin_positions:\t{:}\n\tQ_end_positions:\t{:}".format(Qbeg_pos,Qend_pos))

	# convert to Blackboard import file format
	if verbose:
		print ('\nFound {:} possible question(s), identifying type...'.format(len(Qbeg_pos)))
	for i in range(len(Qbeg_pos)):
		make_Q(i, Qbeg_pos[i], Qend_pos[i])

	# write to Blackboard text file
	with open(WordFileName.replace('.docx','.txt'),'w') as outputfile:
		outputfile.write(BBtext.strip('\n'))
		outputfile.close()
	
	# print summary
	print ('\nSuccessfully wrote question(s) to [{:}]. Summary:'.format(WordFileName.replace('.docx','.txt')))
	SumQ = 0
	for k in sorted(QuestionTypes):
		if k != 'Warning':
			print ('\t{:8}: {:3}'.format(k,QuestionTypes[k]))
			SumQ += QuestionTypes[k]
	print ('\t~~~~~~~~~~~~~\n\t{:8}: {:3}'.format('Total',SumQ))
	if QuestionTypes['Warning'] >0:
		print ('\t{:8}: {:3}'.format('Warning',QuestionTypes['Warning']))
	
def GetListOutline(p):
	"""get if paragraph is a list adnd if so its outline level"""
	
	lst = False
	lvl = 0
	if hasattr(p.paragraph_format.element.pPr, 'numPr'):
		if hasattr(p.paragraph_format.element.pPr.numPr, 'ilvl'):
			lst = True
			if p.paragraph_format.element.pPr.numPr.ilvl != None:
				lvl = p.paragraph_format.element.pPr.numPr.ilvl.val + 1
			else:
				lvl = 1
	if lst == False and lvl == 0:
		if hasattr(p.style.paragraph_format.element.pPr, 'numPr'):
			if hasattr(p.style.paragraph_format.element.pPr.numPr, 'ilvl'):
				lst = True
				if p.style.paragraph_format.element.pPr.numPr.ilvl != None:
					lvl = p.paragraph_format.element.pPr.numPr.ilvl.val + 1
				else:
					lvl = 1
	return lst, lvl

def make_Q(Qid, start, end):
	"""Convert data to Blackboard import file format"""

	global BBtext
	global QuestionTypes
	
	Qid += 1

	BoldCount = 0	#count number of bold answers. 1 = M/C, 2+ Error
	MAT_start = 0	#Matching answer start position
	if end != start:
		for i in range(start+1,end+1):	# range does not include the end value, so +1 is needed
			if data[i]['allBold']: 
				BoldCount += 1
			if MAT_start == 0 and data[i]['outline'] > data[start+1]['outline']:
				MAT_start = i
 
	# T/F question
	if start == end:
		if data[start]['trueBold'] and data[start]['falseBold']:
			print_Warning ("skipped T/F question with both answeres in bold. (Q#{:})\n\t{:}...".format(Qid,data[start]['text'][:(TextWidth-15)]))
			QuestionTypes['Warning'] += 1
		elif data[start]['trueBold']:
			Qtxt = re.sub('\([ ]*True[ ]*/[ ]*False[ ]*\)','',data[start]['text'],flags=re.IGNORECASE).strip(' ')
			BBtext += "\nTF\t{:}\ttrue".format(Qtxt)
			if verbose: print ('\tQ{:} identified as True/False'.format(Qid))
			QuestionTypes['T/F'] += 1
		elif data[start]['falseBold']:
			Qtxt = re.sub('\([ ]*True[ ]*/[ ]*False[ ]*\)','',data[start]['text'],flags=re.IGNORECASE).strip(' ')
			BBtext += "\nTF\t{:}\tfalse".format(Qtxt)
			if verbose: print ('\tQ{:} identified as True/False'.format(Qid))
			QuestionTypes['T/F'] += 1
		else:
			print_Warning ("skipped T/F question with no answeres in bold. (Q#{:})\n\t{:}...".format(Qid,data[start]['text'][:(TextWidth-15)]))
			QuestionTypes['Warning'] += 1
		return

	# Essay question
	if BoldCount == 0 and start+1 == end:	
			BBtext += "\nESS\t{:}\t{:}".format(data[start]['text'],data[end]['text'])
			if verbose: print ('\tQ{:} identified as Essay'.format(Qid))
			QuestionTypes['ESSAY'] += 1
			return

	# M/C question
	if BoldCount == 1:	
		BBtext += "\nMC\t{:}".format(data[start]['text'])
		for i in range(start+1,end+1): # range does not include the end value, so +1 is needed
			if data[i]['allBold']: 	answer = 'correct'
			else: 					answer = 'incorrect'
			BBtext += "\t{:}\t{:}".format(data[i]['text'],answer)
		if verbose: print ('\tQ{:} identified as Multiple Choice'.format(Qid))	
		QuestionTypes['M/C'] += 1
		return

	# Matching question
	n=0
	if BoldCount == 0 and MAT_start != 0:
		if (end - start)%2 == 0 and MAT_start - start - 1 == end - MAT_start +1 : # equal number of sentences and terms
			BBtext += "\nMAT\t{:}".format(data[start]['text'])
			for i in range(start+1,MAT_start):
				BBtext += "\t{:}\t{:}".format(data[start+1+n]['text'],data[MAT_start+n]['text'])
				n += 1
			if verbose: print ('\tQ{:} identified as Matching'.format(Qid))
			QuestionTypes['MAT'] += 1
		else:
			print_Warning ("skipped matching question with unequal count of sentances and terms. (Q#{:})\n\tterms:{:}, sentances:{:}\n\t{:}...".format(Qid,MAT_start-start-1,end-MAT_start+1,data[start]['text'][:(TextWidth-15)]))
			QuestionTypes['Warning'] += 1
		return

	# Fill In the Blank question
	if BoldCount == 0 and re.search('_{5,}',data[start]['text']) != None:	
		BBtext += "\nFIB\t{:}".format(data[start]['text'])
		for i in range(start+1,end+1): # range does not include the end value, so +1 is needed
			BBtext += "\t{:}".format(data[i]['text'])
		if verbose: print ('\tQ{:} identified as Fill_In_the_Blank'.format(Qid))	
		QuestionTypes['FIB'] += 1
		return

	print_Warning ("couldn't identify question type, skipping: (Q#{:}, best guess: ESS or M/C)\n\t{:}...".format(Qid,data[start]['text'][:(TextWidth-15)]))
	QuestionTypes['Warning'] += 1
	
def u2a(txt):
	"""Convert unicode text to ascii"""
	
	val = txt
	printed = False
	for k,v in unicode2ascii["rules"].items():
		if txt.find(k) != -1:
			if verbose:
				if not printed:
					printed = True
					print_to_console ("\t{:}...".format(txt[:(TextWidth-15)]))
				print_to_console ("\t\tconverted {:} to {:}".format(k,v))
			val = val.replace(k,v)
	if verbose:
		found_itr = re.finditer(unicode2ascii["notallowed"],val)
		found_pos = [m.start()+1 for m in found_itr]
		found_val = [txt[m-1] for m in found_pos]
		if found_pos != []:
			if not printed:
				print_to_console ("\t{:}...".format(txt[:(TextWidth-15)]))
			print_Warning ("found {:} unhandled unicode at {:}".format(found_val,found_pos)) 
	return val

### print to console string with any kind of encoding ####################################

def print_Warning(msg):
	"""Print warning message with color"""
	print_to_console (bcolors.WARNING + 'Warning - ' + msg + bcolors.ENDC)

def print_Fail(msg):
	"""Print fail message with color and exit"""
	print_to_console (bcolors.FAIL + 'Error - ' + msg + bcolors.ENDC)
	sys.exit(0)

def print_to_console(text):
	"""Prints a (unicode) string to the console, encoded depending on the stdout encoding 
	(eg. cp437 on Windows). Works with Python 2 and 3."""	
	try:
		sys.stdout.write(text)
	except UnicodeEncodeError:
		bytes_string = text.encode(sys.stdout.encoding, 'backslashreplace')
		if hasattr(sys.stdout, 'buffer'):
			sys.stdout.buffer.write(bytes_string)
		else:
			text = bytes_string.decode(sys.stdout.encoding, 'strict')
			sys.stdout.write(text)
	sys.stdout.write("\n")
    
### Main #################################################################################

if __name__ == "__main__":
	ProcessCLI()
	RunScript()
